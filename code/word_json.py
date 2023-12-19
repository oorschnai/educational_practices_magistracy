import docx
from collections import OrderedDict

"""
Что такое OrderedDict?

Если мы используем set(), то данные будут смещиваться
Например:
lst = [1,2,3]
set(lst)
Out = [2,1,3]

Поэтому мы применяем OrderedDict, т.к он помогает оставить тот же порядок.

* Далее в коментах я буду писать set, но имею ввиду OrderedDict
"""

class Doc_Dict:
    def __init__(self, doc_path):
        self.docx_path = doc_path
        self.doc = docx.Document(self.docx_path.replace("\\", "/"))
        self.list_with_table_values = []

    def create_dict(self,table):
        set_list = []
        for document in range(1, len(table)): # Цикл по спискам, которые не являются "главными"
            s = {} 
            for main in range(len(table[0])): # Цикл по главному списку
                s[table[0][main]] = table[document][main]
            set_list.append(s)
        
        return set_list

    def create_table(self):  

        for table in self.doc.tables: # Цикл по всем таблицам в документе
            
            for current_doc in range(0,len(self.doc.tables)): # Цикл по текущей таблице
                document_set = [] # Список для данных подвергшиеся set()
                document_list = [] # Просто список без изменений
                
                for  en, row in enumerate(table.rows): # Цикл по строке
                    text = [cell.text for cell in row.cells] # Извлекаем данные из ячеек в таблице
                    if en == 0: # условие если это первая строка (ее мы будем использовать как главную)
                        main_text = list(OrderedDict.fromkeys(text)) # Убираем повторы
                        document_set.append(main_text) 
                        continue
                    
                    if len(main_text) == 1: # Если в main_text всего 1 слово, то мы просто оставим все в списке, иначе в dict()
                        document_list.append(text)
                    else:
                        document_set.append(list(OrderedDict.fromkeys(text))) # Убираем повторы из всех ячеек

            if document_list: # Если у нас есть таблица в которой 1 main_text, то без изменений ее добавляем в итоговую
                self.list_with_table_values.append(document_list)
            else: # Иначе мы кидаем все в create_dict и получим словарь
                self.list_with_table_values.append(self.create_dict(table = document_set))

        return self.list_with_table_values


if __name__ == "__main__":
    tables = Doc_Dict("test_files/PRAVILA_OFORMLENIYa_KR_KP_i_VKR_2.docx").create_table()
    for table in tables:
        print("Out = ",table)
        print()